from docx import Document
from docx2pdf import convert
from pdf2docx import Converter

from src.config.settings import RESEARCH_DOCX_PATH, RESEARCH_PAPER_DOCX_PATH, RPAPER_PATH


def pdf_to_docx(pdf_file=None, docx_file=None):
    source_pdf = str(pdf_file or RPAPER_PATH)
    target_docx = str(docx_file or RESEARCH_PAPER_DOCX_PATH)
    cv = Converter(source_pdf)
    cv.convert(target_docx)
    cv.close()


def change_to_one_column(file_path=None):
    source_path = str(file_path or RESEARCH_PAPER_DOCX_PATH)
    doc = Document(source_path)
    new_doc = Document()

    for paragraph in doc.paragraphs:
        new_paragraph = new_doc.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.color.rgb = run.font.color.rgb
            new_run.font.size = run.font.size

    new_doc.save(str(RESEARCH_DOCX_PATH))


def convert_docx_to_pdf():
    convert(str(RESEARCH_DOCX_PATH))
